import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import numpy as np

# Inputs do usuário
p1_ini = float(input("Digite o valor inicial do Parâmetro 1: "))
p1_fim = float(input("Digite o valor final do Parâmetro 1: "))
p2_ini = float(input("Digite o valor inicial do Parâmetro 2: "))
p2_fim = float(input("Digite o valor final do Parâmetro 2: "))

# Carregar os dados
df = pd.read_excel('teste.xlsx')
valores = df['Valor'].tolist()
n = len(valores)

# Cria os parâmetros interpolando do inicial até o final
parametro1 = np.linspace(p1_ini, p1_fim, n)
parametro2 = np.linspace(p2_ini, p2_fim, n)

# Verificar erros
erros = []
for i in range(n):
    if not parametro1[i] <= valores[i] <= parametro2[i]:
        erros.append(f"Linha {i+1}: Valor {valores[i]} fora do intervalo ({parametro1[i]:.2f} - {parametro2[i]:.2f})")

# Plotar gráfico
x = range(len(valores))
plt.figure(figsize=(10, 6))
plt.plot(x, parametro2, color='red', label='Parametro2 (linha superior)')
plt.plot(x, parametro1, color='blue', label='Parametro1 (linha inferior)')
plt.scatter(x, valores, color='black', label='Valor (meio)', zorder=5)
plt.title('Gráfico com verificação de parâmetros')
plt.xlabel('Índice')
plt.ylabel('Valores')
plt.grid(True)
plt.legend()
plt.tight_layout()
plt.savefig('grafico.png')  # Salvar gráfico como imagem
plt.close()

# Criar documento Word
doc = Document()
doc.add_heading('Relatório de Verificação de Parâmetros', level=1)

if erros:
    doc.add_paragraph("🚫 Foram encontrados erros nos valores:")
    for erro in erros:
        doc.add_paragraph(erro, style='List Bullet')
else:
    doc.add_paragraph("✅ Todos os valores estão dentro dos parâmetros!")

# Inserir gráfico no Word
doc.add_picture('grafico.png', width=Inches(5.5))

# Salvar documento
doc.save('relatorio_parametros.docx')
print("📄 Relatório gerado com sucesso: relatorio_parametros.docx")