import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import smtplib
from email.message import EmailMessage
from datetime import datetime

# === 1. INPUTS DO USUÁRIO ===
parametro1_inicial = int(input("Digite o valor inicial de Parametro1: "))
parametro2_inicial = int(input("Digite o valor inicial de Parametro2: "))

# === 2. LER PLANILHA E FAZER ANÁLISE ===
df = pd.read_excel('teste.xlsx')
valores = df['Valor'].tolist()
parametro1 = [parametro1_inicial + i for i in range(len(valores))]
parametro2 = [parametro2_inicial + i for i in range(len(valores))]

erros = []
for i in range(len(valores)):
    if not parametro1[i] <= valores[i] <= parametro2[i]:
        erros.append(f"Linha {i+1}: Valor {valores[i]} fora do intervalo ({parametro1[i]} - {parametro2[i]})")

# === 3. GERAR O GRÁFICO ===
x = range(len(valores))
plt.figure(figsize=(10, 6))
plt.plot(x, parametro2, color='red', label='Parametro2 (linha superior)')
plt.plot(x, parametro1, color='blue', label='Parametro1 (linha inferior)')
plt.scatter(x, valores, color='black', label='Valor (meio)', zorder=5)
plt.title('Gráfico com verificação de parâmetros')
plt.xlabel('Índice')
plt.ylabel('Valores')
plt.grid(True)
plt.legend()
plt.tight_layout()
plt.savefig('grafico.png')
plt.close()

# === 4. CRIAR RELATÓRIO WORD ===
doc = Document()
doc.add_heading('Relatório de Verificação de Parâmetros', level=1)
doc.add_paragraph(f"Data do relatório: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n")

if erros:
    doc.add_paragraph("🚫 Foram encontrados erros nos valores:")
    for erro in erros:
        doc.add_paragraph(erro, style='List Bullet')
else:
    doc.add_paragraph("✅ Todos os valores estão dentro dos parâmetros!")

doc.add_picture('grafico.png', width=Inches(5.5))
doc.save('relatorio_parametros.docx')

# === 5. ENVIAR POR E-MAIL ===
EMAIL_EMISSOR = input("Digite o email emissor: ")
SENHA = input("Digite a senha do email emissor: ")
EMAIL_DESTINO = input("Digite o email destino: ")

# Preparar corpo do e-mail
if erros:
    status_mensagem = "🚫 Foram encontrados erros nos valores:\n"
    for erro in erros:
        status_mensagem += f"- {erro}\n"
else:
    status_mensagem = "✅ Todos os valores estão dentro dos parâmetros!"

msg = EmailMessage()
msg['Subject'] = '📊 Relatório de Parâmetros'
msg['From'] = EMAIL_EMISSOR
msg['To'] = EMAIL_DESTINO
msg.set_content(f"""
Olá,

Segue em anexo o relatório automático gerado pelo sistema Python.

Resumo da verificação:
{status_mensagem}

Atenciosamente,
Sistema de Análise
""")

with open('relatorio_parametros.docx', 'rb') as file:
    msg.add_attachment(file.read(),
                       maintype='application',
                       subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                       filename='relatorio_parametros.docx')

with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
    smtp.login(EMAIL_EMISSOR, SENHA)
    smtp.send_message(msg)

print("📤 Relatório analisado, gerado e enviado por e-mail com sucesso!")
